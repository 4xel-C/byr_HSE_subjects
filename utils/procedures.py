import os
from dataclasses import dataclass
from datetime import datetime

from docx import Document
from docx.shared import Pt
from openpyxl import load_workbook

from .config import HISTORY, MONTH, PATH, PROCEDURES, TEMPLATES


# Data class to represent each procedure
@dataclass
class Procedure:
    number: str  # Number of the procedure, format: PROCXXX
    part: int  # Number of the part (If a long procedure is partionnated)
    title: str
    ignored: bool  # Check if the procedure has to be ignored.
    last_review: datetime = datetime.min  # Date on which the procedure has been reviewed. Completed during the initialization of the procedureMager using the generatre_history() method.


# Class to manage the procedures
class ProcedureManager:
    path = PATH
    procedures_list_file = PROCEDURES
    history_file = HISTORY
    templates_files = TEMPLATES

    def __init__(self) -> None:
        self.procedures = []
        self.__load_procedures()
        self.__generate_history()

    def __load_procedures(self) -> None:
        """Load all the procedures details from a file as a Procedure object and store them in self.procedures"""

        procedures_path = os.path.join(self.path, self.procedures_list_file)

        wb = load_workbook(procedures_path)
        sheet = wb.active

        for number, part, title, ignored in sheet.iter_rows(
            min_row=2, values_only=True
        ):
            self.procedures.append(
                Procedure(number.strip(), part, title.strip(), ignored)
            )
        return

    def __generate_history(self) -> None:
        """Load the history file and get the last reviewed date for each procedure"""

        history_path = os.path.join(self.path, self.history_file)

        wb = load_workbook(history_path)
        sheet = wb.active

        for number, date in sheet.iter_rows(min_row=2, values_only=True):
            
            if not date:
                continue
            
            # check if the date is a datetime object, convert it if needed
            if not isinstance(date, datetime):
                date = datetime.strptime(date, "%m/%Y")
            
            for proc in self.procedures:
                if number.strip() == proc.number:
                    proc.last_review = max(proc.last_review, date)
        return

    def get_procedures(self, n=None) -> list[Procedure]:
        """Select n procedures sorted by date from the oldest to the newest.
        If no n is passed through the function, return all the procedures"""
        procedures = sorted(self.procedures, key=lambda x: x.last_review)

        # return the n first procedures
        return procedures[:n] if n else procedures

    @classmethod
    def write_document(cls, procedures: list[Procedure], quarter: int) -> bool:
        """class method to generate the docx file used for the security discussions based on the selected procedures and the quarter.

        Args:
            procedures (list[Procedure]): A list of Procedure object to consider while creating the file.
            quarter (int): The quarter concerned by the redaction of the file.

        Returns:
            bool: Confirmation of the creation of the file.
        """
        year = datetime.now().year

        # get the template correct template path (2 or 3 rows)
        if quarter == 3:
            template_path = os.path.join(cls.path, cls.templates_files[1])
        else:
            template_path = os.path.join(cls.path, cls.templates_files[0])

        # Load the document and select the table
        doc = Document(template_path)
        table = doc.tables[0]

        # Iterate through all row and the two first columns
        for i in range(len(MONTH[quarter])):
            # select the cells to introduce the informations
            cell_date = table.cell(i + 1, 0)
            cell_subject = table.cell(i + 1, 1)

            # write the text in each cells
            cell_date.text = MONTH[quarter][i]
            cell_subject.text = f"{procedures[i].number}: {procedures[i].title}"

            # update the font of each cells
            for j in range(2):
                cell = table.cell(i + 1, j)
                run = cell.paragraphs[0].runs[0]
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(12)
                cell.paragraphs[0].alignment = 1

        # update the header year
        header = doc.sections[0].header
        paragraph = header.paragraphs[2]
        run = paragraph.add_run(f"{year}")
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(16)

        # Save the document
        try:
            doc.save(
                f"data/Tableau de documentation des discussions mensuelles de sécurité {year} Q{quarter}.docx"
            )
        except Exception:
            return False

        return True

    @classmethod
    def update_history_file(cls, procedures: list[Procedure], quarter: int):
        """Update the history file with the selected procedures for the concerned quarter to keep track of the review for the next iteration.

        Args:
            procedures (list[Procedure]): A list of Procedure object to consider while creating the file.
            quarter (int): The quarter concerned by the redaction of the file.

        Returns:
            bool: Confirmation of the creation of the file.
        """

        try:
            wb = load_workbook(os.path.join(cls.path, cls.history_file))
        except Exception as e:
            print(f"Error while loading the file: {e}")
            return False

        ws = wb.active

        # get the year
        year = str(datetime.today().year)

        # add the new lines
        for i, procedure in enumerate(procedures):
            # Get the month number depending of the quarter considered
            month = (quarter - 1) * 3 + (i + 1)

            # Because July and August are merged together, jump July and go directly to September
            if month == 8:
                month = 9

            month = str(month).zfill(2)
            
            # Parse the string
            date = f"{month}/{year}"
            date = datetime.strptime(date, "%m/%Y")

            ws.append([procedure.number, date])

        try:
            wb.save(os.path.join(cls.path, cls.history_file))
        except Exception as e:
            print(f"Error while saving the file: {e}")
            return False

        return True
